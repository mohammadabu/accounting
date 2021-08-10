# -*- coding: utf-8 -*-

from num2words import num2words
from sys import platform
from subprocess import Popen
from docx2pdf import convert
from datetime import datetime, date
import os
import base64
import logging
from odoo import models, tools,_
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor, Inches , Cm
from docx.oxml.ns import qn,nsdecls
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.dml.color import ColorFormat
from docx.enum.table import WD_TABLE_DIRECTION,WD_ALIGN_VERTICAL,WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement,parse_xml
import docx
_logger = logging.getLogger(__name__)

try:
    from docxtpl import DocxTemplate
except ImportError:
    _logger.debug('Can not import DocxTemplate')

class CustodianReceiptExportDocx(models.AbstractModel):
    _name = 'report.custody_export_docx.custodian_receipt_docx'
    _inherit = 'export.docx.abstract'
    _description = 'Custodian Receipt Export Docx'



    def generate_docx_report(self, data, objs):
        timestamp = str(int(datetime.timestamp(datetime.now())))
        path_docx = '/var/lib/odoo/.local/share/Odoo/'
        custody_id = objs.id
        custody_data = self.pool.get("report.custody_export_docx.custodian_receipt_docx").generate_variables(self,custody_id)
        _logger.info("custody_data")
        _logger.info(custody_data)
        document = docx.Document()
        font_headerTable = document.styles.add_style('font_headerTable', WD_STYLE_TYPE.CHARACTER)
        font_headerTable.font.rtl = True
        font_headerTable.font.size = Pt(16)
        font_headerTable.font.bold = True
        font_headerTable.font.name = 'Arial'

        font_headerTable_1 = document.styles.add_style('font_headerTable_1', WD_STYLE_TYPE.CHARACTER)
        font_headerTable_1.font.rtl = True
        font_headerTable_1.font.size = Pt(12)
        font_headerTable_1.font.bold = True
        font_headerTable_1.font.name = 'Arial' 

        font_headerTable_2 = document.styles.add_style('font_headerTable_2', WD_STYLE_TYPE.CHARACTER)
        font_headerTable_2.font.rtl = True
        font_headerTable_2.font.size = Pt(10)
        font_headerTable_2.font.name = 'Arial' 

        font_headerTable_3 = document.styles.add_style('font_headerTable_3', WD_STYLE_TYPE.CHARACTER)
        font_headerTable_3.font.rtl = True
        font_headerTable_3.font.size = Pt(10)
        font_headerTable_3.font.bold = True
        font_headerTable_3.font.name = 'Arial' 

        # start Header Table  
        headerTable = document.add_table(rows=2,cols=2)

        headerTable.direction = WD_TABLE_DIRECTION.RTL
        headerTable_a_1 = headerTable.cell(0, 0)
        headerTable_b_1 = headerTable.cell(0, 1) 
        headerTable_a_1.merge(headerTable_b_1) 

        headerTable_cells = headerTable.rows[0].cells
        headerTable_cells[0].text = 'نموذج استلام عهده'
        headerTable_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph_headerTable = headerTable_cells[0].paragraphs[0]
        run_headerTable = paragraph_headerTable.runs
        run_headerTable[0].style = font_headerTable
        paragraph_headerTable.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
    
        headerTable_cells_1 = headerTable.rows[1].cells

        headerTable_cells_1[0].text = "التاريخ : ١ / ١ / ٢٠٢١ م"
        paragraph_headerTable_1 = headerTable_cells_1[0].paragraphs[0]
        run_headerTable_1 = paragraph_headerTable_1.runs
        run_headerTable_1[0].style = font_headerTable_1
        paragraph_headerTable_1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        headerTable_cells_1[1].text = "التاريخ : ١/١/١٤٤٢ هـ"
        paragraph_headerTable_2 = headerTable_cells_1[1].paragraphs[0]
        run_headerTable_2 = paragraph_headerTable_2.runs
        run_headerTable_2[0].style = font_headerTable_1
        paragraph_headerTable_2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # End Header Table 

        menuTable = document.add_table(rows=2,cols=6)
        menuTable.direction = WD_TABLE_DIRECTION.RTL
        menuTable.autofit = False 
        menuTable.allow_autofit = False
        menuTable.alignment = WD_TABLE_ALIGNMENT.CENTER

        a = menuTable.cell(0, 0)
        b = menuTable.cell(0,5)
        A = a.merge(b)        
        menuTable.style = 'Table Grid'

        hdr_cells = menuTable.rows[0].cells

        hdr_cells[0].text = 'بيانات المستلم'
        hdr_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        shading_elm_2 = parse_xml(r'<w:shd {} w:fill="#e5e5e5"/>'.format(nsdecls('w')))
        hdr_cells[0]._tc.get_or_add_tcPr().append(shading_elm_2)

        paragraph_header = hdr_cells[0].paragraphs[0]
        run_header = paragraph_header.runs
        run_header[0].style = font_headerTable_1
        paragraph_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


        hdr_cells_item = menuTable.rows[1].cells

        hdr_cells_item[0].text = str(custody_data.job_id)
        hdr_cells_item[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph_item_0 = hdr_cells_item[0].paragraphs[0]
        run_item_0 = paragraph_item_0.runs
        run_item_0[0].style = font_headerTable_2
        paragraph_item_0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        hdr_cells_item[1].text = 'المسمى الوظيفي'

        hdr_cells_item[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shading_cells_item_1 = parse_xml(r'<w:shd {} w:fill="#e5e5e5"/>'.format(nsdecls('w')))
        hdr_cells_item[1]._tc.get_or_add_tcPr().append(shading_cells_item_1)
        paragraph_item_1 = hdr_cells_item[1].paragraphs[0]
        run_item_1 = paragraph_item_1.runs
        run_item_1[0].style = font_headerTable_3
        paragraph_item_1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER



        hdr_cells_item[2].text = custody_data.department
        hdr_cells_item[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph_item_2 = hdr_cells_item[2].paragraphs[0]
        run_item_2 = paragraph_item_2.runs
        run_item_2[0].style = font_headerTable_2
        paragraph_item_2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


        hdr_cells_item[3].text = 'الإدارة'
        hdr_cells_item[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shading_cells_item_3 = parse_xml(r'<w:shd {} w:fill="#e5e5e5"/>'.format(nsdecls('w')))
        hdr_cells_item[3]._tc.get_or_add_tcPr().append(shading_cells_item_3)
        paragraph_item_3 = hdr_cells_item[3].paragraphs[0]
        run_item_3 = paragraph_item_3.runs
        run_item_3[0].style = font_headerTable_3
        paragraph_item_3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


        hdr_cells_item[4].text = custody_data.employee_name
        hdr_cells_item[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph_item_4 = hdr_cells_item[4].paragraphs[0]
        run_item_4 = paragraph_item_4.runs
        run_item_4[0].style = font_headerTable_2
        paragraph_item_4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


        hdr_cells_item[5].text = 'اسـم المـوظف'
        hdr_cells_item[5].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shading_cells_item = parse_xml(r'<w:shd {} w:fill="#e5e5e5"/>'.format(nsdecls('w')))
        hdr_cells_item[5]._tc.get_or_add_tcPr().append(shading_cells_item)
        paragraph_item_5 = hdr_cells_item[5].paragraphs[0]
        run_item_5 = paragraph_item_5.runs
        run_item_5[0].style = font_headerTable_3
        paragraph_item_5.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


        for cell in menuTable.columns[0].cells:
            cell.width = Cm(4.8)
        for cell in menuTable.columns[1].cells:
            cell.width = Cm(2.7)
        for cell in menuTable.columns[2].cells:
            cell.width = Cm(4.1)        
        for cell in menuTable.columns[3].cells:
            cell.width = Cm(1.3)
        for cell in menuTable.columns[4].cells:
            cell.width = Cm(3.2)
        for cell in menuTable.columns[5].cells:
            cell.width = Cm(2.1)            


        count = 1
        for row in menuTable.rows:
            if count == 2: 
                row.height = Cm(1.7)  
            else:
                row.height = Cm(0.7)   
            count = count + 1      

        tbl = menuTable._tbl
        count = 1
        for cell in tbl.iter_tcs():
            tcPr = cell.tcPr
            tcBorders = OxmlElement('w:tcBorders')

            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'double')

            if count != 1:
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'double')

            if count != 1 and count != 2:
                left = OxmlElement('w:left')
                left.set(qn('w:val'), 'single')
            else:
                left = OxmlElement('w:left')
                left.set(qn('w:val'), 'double')

            if count != 1 and count != 7:
                right = OxmlElement('w:right')
                right.set(qn('w:val'), 'single')
            else :
                right = OxmlElement('w:right')
                right.set(qn('w:val'), 'double')    

                

            tcBorders.append(top)
            tcBorders.append(left)
            if count != 1:
                tcBorders.append(bottom)
            tcBorders.append(right)
            tcPr.append(tcBorders)
            count = count + 1

        paragraph = document.add_paragraph('')   
        paragraph_format = paragraph.paragraph_format
        paragraph_format.line_spacing = 0.06

        #table 2
        record = [
            ['2','6','[300030001] آلة تصوير Canon ir601n','[300030001] آلة تصوير Canon ir601n','1'],
            ['2','6','[300030001] آلة تصوير Canon ir601n','[300030001] آلة تصوير Canon ir601n','2'],
            ['2','4','[300030001] آلة تصوير Canon ir601n','[300030001] آلة تصوير Canon ir601n','3'],
        ] 

        subTable = document.add_table(rows=2,cols=5)
        subTable.direction = WD_TABLE_DIRECTION.RTL
        subTable.autofit = False 
        subTable.allow_autofit = False
        subTable.alignment = WD_TABLE_ALIGNMENT.CENTER

        a = subTable.cell(0, 0)
        b = subTable.cell(0,4)
        a.merge(b)

        subTable.style = 'Table Grid'

        sub_hdr_cells = subTable.rows[0].cells

        sub_hdr_cells[0].text = 'بيانات العهدة'
        sub_hdr_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        sub_shading_elm_2 = parse_xml(r'<w:shd {} w:fill="#e5e5e5"/>'.format(nsdecls('w')))
        sub_hdr_cells[0]._tc.get_or_add_tcPr().append(sub_shading_elm_2)
        sub_paragraph_header = sub_hdr_cells[0].paragraphs[0]
        sub_run_header = sub_paragraph_header.runs
        sub_run_header[0].style = font_headerTable_1
        sub_paragraph_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        sub_hdr_cells_item = subTable.rows[1].cells

        sub_hdr_cells_item[0].text = 'ملاحظة'
        sub_hdr_cells_item[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        sub_paragraph_item_0 = sub_hdr_cells_item[0].paragraphs[0]
        sub_run_item_0 = sub_paragraph_item_0.runs
        sub_run_item_0[0].style = font_headerTable_2
        sub_paragraph_item_0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        sub_hdr_cells_item[1].text = 'الكمية'
        sub_hdr_cells_item[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        sub_paragraph_item_1 = sub_hdr_cells_item[1].paragraphs[0]
        sub_run_item_1 = sub_paragraph_item_1.runs
        sub_run_item_1[0].style = font_headerTable_2
        sub_paragraph_item_1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        sub_hdr_cells_item[2].text = 'النوع'
        sub_hdr_cells_item[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        sub_paragraph_item_2 = sub_hdr_cells_item[2].paragraphs[0]
        sub_run_item_2 = sub_paragraph_item_2.runs
        sub_run_item_2[0].style = font_headerTable_2
        sub_paragraph_item_2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        sub_hdr_cells_item[3].text = 'الوصف'
        sub_hdr_cells_item[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        sub_paragraph_item_3 = sub_hdr_cells_item[3].paragraphs[0]
        sub_run_item_3 = sub_paragraph_item_3.runs
        sub_run_item_3[0].style = font_headerTable_2
        sub_paragraph_item_3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        sub_hdr_cells_item[4].text = 'م'
        sub_hdr_cells_item[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        sub_paragraph_item_4 = sub_hdr_cells_item[4].paragraphs[0]
        sub_run_item_4 = sub_paragraph_item_4.runs
        sub_run_item_4[0].style = font_headerTable_2
        sub_paragraph_item_4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


        # dynamic

        for ID, name, price,tt,ttt in record:
            row_Cells_sub_dynamic = subTable.add_row().cells

            row_Cells_sub_dynamic[0].text = ID

            row_Cells_sub_dynamic[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            sub_dynamic_paragraph_0 = row_Cells_sub_dynamic[0].paragraphs[0]
            sub_dynamic_run_0 = sub_dynamic_paragraph_0.runs
            sub_dynamic_run_0[0].style = font_headerTable_2
            sub_dynamic_paragraph_0.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


            row_Cells_sub_dynamic[1].text = name

            row_Cells_sub_dynamic[1].text = ID
            row_Cells_sub_dynamic[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            sub_dynamic_paragraph_1 = row_Cells_sub_dynamic[1].paragraphs[0]
            sub_dynamic_run_1 = sub_dynamic_paragraph_1.runs
            sub_dynamic_run_1[0].style = font_headerTable_2
            sub_dynamic_paragraph_1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            row_Cells_sub_dynamic[2].text = price

            row_Cells_sub_dynamic[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            sub_dynamic_paragraph_2 = row_Cells_sub_dynamic[2].paragraphs[0]
            sub_dynamic_run_2 = sub_dynamic_paragraph_2.runs
            sub_dynamic_run_2[0].style = font_headerTable_2
            sub_dynamic_paragraph_2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


            row_Cells_sub_dynamic[3].text = tt

            row_Cells_sub_dynamic[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            sub_dynamic_paragraph_3 = row_Cells_sub_dynamic[3].paragraphs[0]
            sub_dynamic_run_3 = sub_dynamic_paragraph_3.runs
            sub_dynamic_run_3[0].style = font_headerTable_2
            sub_dynamic_paragraph_3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


            row_Cells_sub_dynamic[4].text = ttt

            row_Cells_sub_dynamic[4].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            sub_dynamic_paragraph_4 = row_Cells_sub_dynamic[4].paragraphs[0]
            sub_dynamic_run_4 = sub_dynamic_paragraph_4.runs
            sub_dynamic_run_4[0].style = font_headerTable_2
            sub_dynamic_paragraph_4.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # end dynmic

        # start gray cell 
        row_Cells_sub_gray = subTable.add_row().cells
        row_Cells_sub_gray[0].text = ''
        gray_shading_elm_0 = parse_xml(r'<w:shd {} w:fill="#bfbfbf"/>'.format(nsdecls('w')))
        row_Cells_sub_gray[0]._tc.get_or_add_tcPr().append(gray_shading_elm_0)
        row_Cells_sub_gray[1].text = ''
        gray_shading_elm_1 = parse_xml(r'<w:shd {} w:fill="#bfbfbf"/>'.format(nsdecls('w')))
        row_Cells_sub_gray[1]._tc.get_or_add_tcPr().append(gray_shading_elm_1)
        row_Cells_sub_gray[2].text = ''
        gray_shading_elm_2 = parse_xml(r'<w:shd {} w:fill="#bfbfbf"/>'.format(nsdecls('w')))
        row_Cells_sub_gray[2]._tc.get_or_add_tcPr().append(gray_shading_elm_2)
        row_Cells_sub_gray[3].text = ''
        gray_shading_elm_3 = parse_xml(r'<w:shd {} w:fill="#bfbfbf"/>'.format(nsdecls('w')))
        row_Cells_sub_gray[3]._tc.get_or_add_tcPr().append(gray_shading_elm_3)
        row_Cells_sub_gray[4].text = ''
        gray_shading_elm_4 = parse_xml(r'<w:shd {} w:fill="#bfbfbf"/>'.format(nsdecls('w')))
        row_Cells_sub_gray[4]._tc.get_or_add_tcPr().append(gray_shading_elm_4)
        # end gray cell 

        # start akrar cell
        row_Cells_sub_akrar = subTable.add_row().cells
        akrar_a = subTable.cell(3 + len(record), 0)
        akrar_b = subTable.cell(3 + len(record), 4) 
        akrar_a.merge(akrar_b)  

        row_Cells_sub_akrar[0].text = 'إقـرار'
        row_Cells_sub_akrar[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        akrar_shading_elm_0 = parse_xml(r'<w:shd {} w:fill="#e5e5e5"/>'.format(nsdecls('w')))
        row_Cells_sub_akrar[0]._tc.get_or_add_tcPr().append(akrar_shading_elm_0)
        akrar_paragraph_header = row_Cells_sub_akrar[0].paragraphs[0]
        akrar_run_header = akrar_paragraph_header.runs
        akrar_run_header[0].style = font_headerTable_1
        akrar_paragraph_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        row_Cells_sub_akrar_1 = subTable.add_row().cells
        date_rec = "23/05/2021 م  الموافق  الاحد "
        day_rec = "الاحد"
        paragraph_marge = ('أقر أنا الموقع أدناه بأنني استلمت العُهد الموضحة أعلاه في يوم %s بحالة جديدة، ') % (date_rec)
        p = row_Cells_sub_akrar_1[0].add_paragraph(paragraph_marge)
        run_manager = p.runs
        font_manager = run_manager[0].font
        font_manager.size = Pt(10)
        font_manager.bold = True
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        runner = p.add_run("وأتعهد بالمحافظة عليها وان لا أتنازل عنها لأي شخص آخر وسأقوم بإعادتها عند طلبها أو عند ترك العمل أو دفع قيمة ما تسببت في تلفه وسأكون عرضة للمسائلة في حين مخالفتي للإقرار.")
        runner.bold = True
        font = runner.font
        font.color.rgb = RGBColor(255,0,0)
        runner.style = font_headerTable_3
        
        
        p1 = row_Cells_sub_akrar_1[0].add_paragraph("\n")

        run_manager_p1 = p1.runs
        font_manager_p1 = run_manager_p1[0].font
        font_manager_p1.size = Pt(10)
        font_manager_p1.bold = True
        p1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        runner_1 = p1.add_run("..................................................... / المستلم")
        runner_1 = p1.add_run("..................................................... / التوقيع")
        runner_1.style = font_headerTable_3
        # runner.bold = True
        # font = runner.font
        # font.color.rgb = RGBColor(255,0,0)
        # runner.style = font_headerTable_3



        
        akrar_a_1 = subTable.cell(4 + len(record), 0)
        akrar_b_1 = subTable.cell(4 + len(record), 4) 
        akrar_a_1.merge(akrar_b_1)  




        # end akrar cell

        # 18.2
        for cell in subTable.columns[0].cells:
            cell.width = Cm(6)
        for cell in subTable.columns[1].cells:
            cell.width = Cm(2.6)
        for cell in subTable.columns[2].cells:
            cell.width = Cm(2.6)        
        for cell in subTable.columns[3].cells:
            cell.width = Cm(6)
        for cell in subTable.columns[4].cells:
            cell.width = Cm(1)  

        count = 1
        for row in subTable.rows:
            if count == 1: 
                row.height = Cm(0.8)  
            elif 5 + len(record) == count:
                row.height = Cm(3)  
            else:
                row.height = Cm(0.7)   
            count = count + 1





        paragraph_last = document.add_paragraph('')   
        last_paragraph_format = paragraph_last.paragraph_format
        last_paragraph_format.line_spacing = 0.06

        #last table 

        lastTable = document.add_table(rows=3,cols=4)
        lastTable.direction = WD_TABLE_DIRECTION.RTL
        lastTable.autofit = False 
        lastTable.allow_autofit = False
        lastTable.alignment = WD_TABLE_ALIGNMENT.CENTER

        a_lastTable = lastTable.cell(0, 0)
        b_lastTable = lastTable.cell(0,3)
        a_lastTable.merge(b_lastTable)   
        a1_lastTable = lastTable.cell(1, 0)
        b1_lastTable = lastTable.cell(1,3)
        a1_lastTable.merge(b1_lastTable)        
        lastTable.style = 'Table Grid'


        last_hdr_cells = lastTable.rows[0].cells

        msg_text = "خاص بمسؤول التسليم والاستلام"
        msg_text = msg_text + ")"
        msg_text = msg_text + " قسم دعم الأعمال  "
        msg_text = msg_text + "(ه"




        # last_hdr_cells[0].text = msg_text
        last_parag_1_1 = last_hdr_cells[0].add_paragraph(" ")
        
        last_parag_1_1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        last_parag_1_1.paragraph_format.line_spacing = 1.5
        last_parag_1_1.paragraph_format.right_indent = Cm(0.6)





        last_1_runner_1 = last_parag_1_1.add_run("خاص بمسؤول التسليم والاستلام")
        last_1_runner_1.style = font_headerTable_1
        last_1_runner_2 = last_parag_1_1.add_run(")")
        last_1_runner_2.style = font_headerTable_1
        last_1_runner_3 = last_parag_1_1.add_run(" قسم دعم الأعمال  ")
        last_1_runner_3.style = font_headerTable_1
        last_1_runner_4 = last_parag_1_1.add_run("(")
        last_1_runner_4.style = font_headerTable_1
        last_1_runner_5 = last_parag_1_1.add_run("ه")
        last_1_runner_5.style = font_headerTable_1
        font_1_runner_5 = last_1_runner_5.font
        font_1_runner_5.color.rgb = RGBColor(229,229,229)
        last_1_runner_1.style = font_headerTable_1

        
        
        last_hdr_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        last_shading_elm_2 = parse_xml(r'<w:shd {} w:fill="#e5e5e5"/>'.format(nsdecls('w')))
        last_hdr_cells[0]._tc.get_or_add_tcPr().append(last_shading_elm_2)
        # last_paragraph_header = last_hdr_cells[0].paragraphs[0]
        # last_run_header = last_paragraph_header.runs
        # last_run_header[0].style = font_headerTable_1
        # last_paragraph_header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        


        last_hdr_cells_1 = lastTable.rows[1].cells

        last_p1 = last_hdr_cells_1[0].add_paragraph(" ")

        last_run_manager_p1 = last_p1.runs
        last_font_manager_p1 = last_run_manager_p1[0].font
        last_p1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        last_p1.paragraph_format.line_spacing = 1.5
        last_p1.paragraph_format.right_indent = Cm(0.6)
        last_runner_1 = last_p1.add_run('أ‌-    تم تسليم العهدة وفقاً للبيان أعلاه بحالة تسليم ')
        font_headerTable_4 = document.styles.add_style('font_headerTable_4', WD_STYLE_TYPE.CHARACTER)
        last_runner_1.style = font_headerTable_3

        last_runner_2 = last_p1.add_run(')')
        last_runner_2.style = font_headerTable_3
        last_runner_3 = last_p1.add_run('جديد')
        last_runner_3.style = font_headerTable_3
        last_runner_4 = last_p1.add_run('(')
        last_runner_4.style = font_headerTable_3
        last_runner_5 = last_p1.add_run('.')
        last_runner_6 = last_p1.add_run('ه')
        font_6 = last_runner_6.font
        font_6.color.rgb = RGBColor(255,255,255)
        last_runner_6.style = font_headerTable_3

        last_p2 = last_hdr_cells_1[0].add_paragraph(" ")
        last_run_manager_p2 = last_p2.runs
        last_font_manager_p2 = last_run_manager_p2[0].font
        last_p2.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        last_p2.paragraph_format.line_spacing = 1.5
        last_p2.paragraph_format.right_indent = Cm(0.6)
        last_runner_2 = last_p1.add_run(' ')
        last_runner_2.style = font_headerTable_3
        runner_p2_1 = last_p2.add_run("........................................./ التسليم بواسطة")
        runner_p2_2 = last_p2.add_run("................................................/ التوقيع")
        runner_p2_1.style = font_headerTable_3
        runner_p2_2.style = font_headerTable_3



        last_p3 = last_hdr_cells_1[0].add_paragraph(" ")

        last_run_manager_p3 = last_p3.runs
        last_font_manager_p3 = last_run_manager_p3[0].font
        last_p3.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        last_p3.paragraph_format.line_spacing = 1.5
        last_p3.paragraph_format.right_indent = Cm(0.6)
        last_runner_3 = last_p3.add_run('ب-    تم استلام العُهد ')
        last_runner_3.style = font_headerTable_3
        last_runner_4 = last_p3.add_run('( ')
        last_runner_4.style = font_headerTable_3
        last_runner_5 = last_p3.add_run(' )')
        last_runner_5.style = font_headerTable_3
        last_runner_6 = last_p3.add_run(' نعم ')
        last_runner_6.style = font_headerTable_3
        last_runner_7 = last_p3.add_run(' سسسسسسسسسسسسسسس ')
        font_7 = last_runner_7.font
        font_7.color.rgb = RGBColor(255,255,255)        
        last_runner_8 = last_p3.add_run(' (  ) لا، للأسباب التالية  ')
        last_runner_8.style = font_headerTable_3
        


        last_p4 = last_hdr_cells_1[0].add_paragraph(" ")
        last_p4.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        last_p4.paragraph_format.line_spacing = 1.5
        last_p4.paragraph_format.right_indent = Cm(0.6)
        last_runner_4 = last_p4.add_run('................................................................................................................    -1')
        last_runner_4.style = font_headerTable_3
        
        last_p5 = last_hdr_cells_1[0].add_paragraph(" ")
        last_p5.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        last_p5.paragraph_format.line_spacing = 1.5
        last_p5.paragraph_format.right_indent = Cm(0.6)
        last_runner_5 = last_p5.add_run('................................................................................................................    -2')
        last_runner_5.style = font_headerTable_3

        last_p6 = last_hdr_cells_1[0].add_paragraph(" ")
        last_p6.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        last_p6.paragraph_format.line_spacing = 1.5
        last_p6.paragraph_format.right_indent = Cm(0.6)
        last_runner_6 = last_p6.add_run('................................................................................................................    -3')
        last_runner_6.style = font_headerTable_3


        last_p7 = last_hdr_cells_1[0].add_paragraph("")
        last_p7.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        last_p7.paragraph_format.line_spacing = 1.5
        last_p7.paragraph_format.right_indent = Cm(0.6)
        last_runner_7 = last_p7.add_run('تم استلام العُهد في يوم/')
        last_runner_7.style = font_headerTable_3
        last_runner_7_1 = last_p7.add_run('سسسسسسسس')
        font_7_1 = last_runner_7_1.font
        font_7_1.color.rgb = RGBColor(255,255,255)   
        last_runner_8 = last_p7.add_run('الموافق ')
        last_runner_8.style = font_headerTable_3

        last_runner_9 = last_p7.add_run('سس')
        font_9 = last_runner_9.font
        font_9.color.rgb = RGBColor(255,255,255)   
        last_runner_9.style = font_headerTable_3

        last_runner_10 = last_p7.add_run('/')
        last_runner_10.style = font_headerTable_3

        last_runner_10 = last_p7.add_run('سس')
        font_10 = last_runner_10.font
        font_10.color.rgb = RGBColor(255,255,255)
        last_runner_10.style = font_headerTable_3

        last_runner_10 = last_p7.add_run('/2020م')
        last_runner_10.style = font_headerTable_3
        
        

        last_hdr_cells_2 = lastTable.rows[2].cells     

        last_hdr_cells_2[0].text = ''
        last_hdr_cells_2[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # shading_cells_item_1 = parse_xml(r'<w:shd {} w:fill="#e5e5e5"/>'.format(nsdecls('w')))
        # last_hdr_cells_2[0]._tc.get_or_add_tcPr().append(shading_cells_item_1)
        paragraph_item_1 = last_hdr_cells_2[0].paragraphs[0]
        run_item_1 = paragraph_item_1.runs
        run_item_1[0].style = font_headerTable_3
        paragraph_item_1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        last_hdr_cells_2[1].text = 'التوقيع'
        last_hdr_cells_2[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # shading_cells_item_1_1 = parse_xml(r'<w:shd {} w:fill="#e5e5e5"/>'.format(nsdecls('w')))
        # last_hdr_cells_2[1]._tc.get_or_add_tcPr().append(shading_cells_item_1_1)
        paragraph_item_1_1 = last_hdr_cells_2[1].paragraphs[0]
        run_item_1_1 = paragraph_item_1_1.runs
        run_item_1_1[0].style = font_headerTable_3
        paragraph_item_1_1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


        last_hdr_cells_2[2].text = 'مشاري بن رجاء المحلسي'
        last_hdr_cells_2[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph_item_1_1 = last_hdr_cells_2[2].paragraphs[0]
        run_item_1_1 = paragraph_item_1_1.runs
        font_headerTable_3_3_3 = document.styles.add_style('font_headerTable_3_3_3', WD_STYLE_TYPE.CHARACTER)
        font_headerTable_3_3_3.font.rtl = True
        font_headerTable_3_3_3.font.size = Pt(9)
        font_headerTable_3_3_3.font.bold = True
        font_headerTable_3_3_3.font.name = 'Arial' 
        run_item_1_1[0].style = font_headerTable_3_3_3
        paragraph_item_1_1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        last_hdr_cells_2[3].text = 'مدير قسم دعم الأعمال'
        last_hdr_cells_2[3].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        paragraph_item_1_1 = last_hdr_cells_2[3].paragraphs[0]
        run_item_1_1 = paragraph_item_1_1.runs
        run_item_1_1[0].style = font_headerTable_3
        paragraph_item_1_1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        
        for cell in lastTable.columns[0].cells: 
            cell.width = Cm(6)
        for cell in lastTable.columns[1].cells:
            cell.width = Cm(2)
        for cell in lastTable.columns[2].cells:
            cell.width = Cm(4.2)        
        for cell in lastTable.columns[3].cells:
            cell.width = Cm(6)



        last_tbl = lastTable._tbl
        count = 1
        for cell in last_tbl.iter_tcs():
            tcPr = cell.tcPr
            tcBorders = OxmlElement('w:tcBorders')

            if count == 1 or count == 2:
                top = OxmlElement('w:top')
                top.set(qn('w:val'), 'double')
            else:
                top = OxmlElement('w:top')
                top.set(qn('w:val'), 'single')

            if count == 3 or count == 4 or count == 5 or count == 6 :
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'double')
            else:
                bottom = OxmlElement('w:bottom')
                bottom.set(qn('w:val'), 'single')    

            if count == 1 or count == 2 or count == 3:
                left = OxmlElement('w:left')
                left.set(qn('w:val'), 'double')
            else:
                left = OxmlElement('w:left')
                left.set(qn('w:val'), 'single')

            if count == 1 or count == 2 or count == 6:
                right = OxmlElement('w:right')
                right.set(qn('w:val'), 'double')   
            else :
                right = OxmlElement('w:right')
                right.set(qn('w:val'), 'single')

            tcBorders.append(top)
            tcBorders.append(left)
            tcBorders.append(bottom)
            tcBorders.append(right)
            tcPr.append(tcBorders)
            count = count + 1  

        
        count = 1
        for row in lastTable.rows:
            if count == 1 or count == 3: 
                row.height = Cm(0.8)  
            else:
                row.height = Cm(5)   
            count = count + 1

            
             
        path_docx = path_docx + '/EmployeeDocx_' + timestamp + "_" + "2131232132131321321" + ".docx"
        document.save(path_docx)
        
        report_doxc_path = path_docx
        with open(report_doxc_path, mode='rb') as file:
            fileContent = file.read()

        try:
            os.remove(report_doxc_path)
        except Exception as e:
            _logger.warning(repr(e))

        return fileContent

    def generate_variables(self, custody_id):
        custody_info = self.env['hr.custody'].sudo().search([('id','=',custody_id)])
        employee_id = custody_info.employee.id
        employee_info = self.env['hr.employee'].sudo().search([('id','=',employee_id)])
        job_id = employee_info.job_id.id
        # job_id = self.env['hr.job'].sudo().search([('id','=',job_id)],limit=1).name
        employee_name = employee_info.name
        department = employee_info.department_id.name
        custody_lines = custody_info.custody_lines

        # not working 
        job_position = self.env['hr.job'].sudo().search([('name','=','مدير قسم دعم الأعمال')],limit=1)
        job_position_id = job_position.id
        employee_info_supp = self.env['hr.employee'].sudo().search([('job_id','=',job_position_id)],limit=1)
        business_support_manager = employee_info_supp.name

        employee_date_array = {}
        employee_date_array['employee_name'] = employee_name
        employee_date_array['department'] = department
        employee_date_array['business_support_manager'] = business_support_manager
        employee_date_array['custody_lines'] = custody_lines
        employee_date_array['job_id'] = job_id
        return employee_date_array            


    def get_report_name(self, objs):
        timestamp = str(int(datetime.timestamp(datetime.now())))
        report_name = f'{objs.sale_order_id.id}_{objs.report_template_id.id}_{timestamp}_report.docx'
        return report_name


